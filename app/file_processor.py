import io
import csv
import base64
from typing import Tuple
from PyPDF2 import PdfReader
from docx import Document
import openpyxl


def process_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF"""
    reader = PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n\n"
    
    if not text.strip():
        return "ERROR: Could not extract text from PDF. The PDF might be image-based."
    
    return text.strip()


def process_excel(file_bytes: bytes) -> str:
    """Extract data from Excel"""
    workbook = openpyxl.load_workbook(io.BytesIO(file_bytes))
    text = ""
    
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        text += f"\n=== Sheet: {sheet_name} ===\n\n"
        
        for row in sheet.iter_rows(values_only=True):
            row_data = [str(cell) if cell is not None else "" for cell in row]
            text += " | ".join(row_data) + "\n"
    
    return text.strip()


def process_word(file_bytes: bytes) -> str:
    """Extract text from Word document"""
    doc = Document(io.BytesIO(file_bytes))
    text = ""
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n\n"
    
    # Also extract tables
    for table in doc.tables:
        text += "\n[TABLE]\n"
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            text += " | ".join(row_data) + "\n"
        text += "[/TABLE]\n\n"
    
    return text.strip()


def process_csv_file(file_bytes: bytes) -> str:
    """Extract data from CSV"""
    content = file_bytes.decode("utf-8")
    reader = csv.reader(io.StringIO(content))
    text = ""
    
    for row in reader:
        text += " | ".join(row) + "\n"
    
    return text.strip()


def process_text(file_bytes: bytes) -> str:
    """Extract text from plain text file"""
    return file_bytes.decode("utf-8").strip()


def process_image(file_bytes: bytes) -> Tuple[str, str]:
    """Convert image to base64 for vision models"""
    base64_image = base64.b64encode(file_bytes).decode("utf-8")
    return base64_image, "image"


def process_file(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """
    Process any file and return (extracted_content, file_type)
    """
    extension = filename.lower().split(".")[-1]
    
    processors = {
        "pdf": ("pdf", lambda b: process_pdf(b)),
        "xlsx": ("excel", lambda b: process_excel(b)),
        "xls": ("excel", lambda b: process_excel(b)),
        "docx": ("word", lambda b: process_word(b)),
        "doc": ("word", lambda b: process_word(b)),
        "csv": ("csv", lambda b: process_csv_file(b)),
        "txt": ("text", lambda b: process_text(b)),
        "md": ("text", lambda b: process_text(b)),
        "png": ("image", None),
        "jpg": ("image", None),
        "jpeg": ("image", None),
        "gif": ("image", None),
        "webp": ("image", None),
    }
    
    if extension not in processors:
        return f"Unsupported file type: .{extension}", "error"
    
    file_type, processor = processors[extension]
    
    if file_type == "image":
        base64_img, _ = process_image(file_bytes)
        return base64_img, "image"
    
    try:
        content = processor(file_bytes)
        # Limit text to avoid token limits
        if len(content) > 8000:
            content = content[:8000] + "\n\n... [Content truncated to fit token limits]"
        return content, file_type
    except Exception as e:
        return f"Error processing file: {str(e)}", "error"


def get_supported_formats() -> dict:
    return {
        "documents": [".pdf", ".docx", ".doc", ".txt", ".md"],
        "spreadsheets": [".xlsx", ".xls", ".csv"],
        "images": [".png", ".jpg", ".jpeg", ".gif", ".webp"]
    }
